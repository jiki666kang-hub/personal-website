from pathlib import Path
import re
from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
RELEASE_ROOT = ROOT.parents[1]
DOCS = RELEASE_ROOT / "assets" / "docs"

MATERIALS = [
    {
        "id": "survey-overseas-resume",
        "title": "大地测量外派俄罗斯方向简历",
        "source": DOCS / "简历" / "汲康康-大地测量外派俄罗斯方向简历-修订版.docx",
        "pdf": ROOT / "pdfs" / "survey-overseas-resume.pdf",
    },
    {
        "id": "survey-resume",
        "title": "大地测量方向简历",
        "source": DOCS / "简历" / "汲康康-大地测量方向简历-修订版.docx",
        "pdf": ROOT / "pdfs" / "survey-resume.pdf",
    },
    {
        "id": "russian-sales-resume",
        "title": "俄语销售方向简历",
        "source": DOCS / "简历" / "汲康康-俄语销售方向简历.docx",
        "pdf": ROOT / "pdfs" / "russian-sales-resume.pdf",
    },
    {
        "id": "interview-cert-handbook",
        "title": "求职面试与留服认证行动手册",
        "source": DOCS / "求职规划" / "汲康康-求职面试与留服认证行动手册.docx",
        "pdf": ROOT / "pdfs" / "interview-cert-handbook.pdf",
    },
    {
        "id": "en-ru-intro-card",
        "title": "英俄语面试自我介绍速记卡",
        "source": DOCS / "英俄面试练习器" / "汲康康-英俄语面试自我介绍速记卡.docx",
        "pdf": ROOT / "pdfs" / "en-ru-intro-card.pdf",
    },
    {
        "id": "en-ru-intro-card-source",
        "title": "英俄语面试自我介绍速记卡（求职规划原件）",
        "source": DOCS / "求职规划" / "汲康康-英俄语面试自我介绍速记卡.docx",
        "pdf": ROOT / "pdfs" / "en-ru-intro-card-source.pdf",
    },
]


def register_font():
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont("MaterialFont", str(path), subfontIndex=0))
                return "MaterialFont"
            except TypeError:
                try:
                    pdfmetrics.registerFont(TTFont("MaterialFont", str(path)))
                    return "MaterialFont"
                except Exception:
                    continue
            except Exception:
                continue
    return "Helvetica"


def escape(text):
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def normalize(text):
    text = re.sub(r"\s+", " ", text.strip())
    return text


def extract_docx(path):
    doc = Document(str(path))
    blocks = []
    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text:
            style = paragraph.style.name if paragraph.style else ""
            blocks.append(("heading" if "Heading" in style or "标题" in style else "para", text))
    for table in doc.tables:
        for row in table.rows:
            cells = [normalize(cell.text.replace("\n", " / ")) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                blocks.append(("table", " | ".join(cells)))
    return blocks


def build_pdf(material, font_name):
    material["pdf"].parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(material["pdf"]),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=material["title"],
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        leading=24,
        textColor=colors.HexColor("#111827"),
        spaceAfter=9,
        alignment=TA_LEFT,
    )
    meta_style = ParagraphStyle(
        "MetaCN",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor("#667085"),
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "HeadingCN",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=12.5,
        leading=17,
        textColor=colors.HexColor("#0f766e"),
        spaceBefore=8,
        spaceAfter=5,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor("#1f2933"),
        spaceAfter=5,
    )
    table_style = ParagraphStyle(
        "TableLineCN",
        parent=body_style,
        backColor=colors.HexColor("#f6f7f9"),
        borderColor=colors.HexColor("#d8dee8"),
        borderWidth=0.4,
        borderPadding=4,
        spaceAfter=5,
    )

    story = [
        Paragraph(escape(material["title"]), title_style),
        Paragraph(f"来源文件：{escape(material['source'].name)}", meta_style),
    ]
    blocks = extract_docx(material["source"])
    for kind, text in blocks:
        if len(story) > 360 and kind == "heading":
            story.append(PageBreak())
        style = table_style if kind == "table" else heading_style if kind == "heading" else body_style
        story.append(Paragraph(escape(text), style))
    story.append(Spacer(1, 8))
    story.append(Paragraph("说明：此 PDF 用于站内预览；如需保留原始 Word 排版，请使用页面中的“打开原始 Word 文件”。", meta_style))
    doc.build(story)


def main():
    font_name = register_font()
    missing = [str(item["source"]) for item in MATERIALS if not item["source"].exists()]
    if missing:
        raise FileNotFoundError("\n".join(missing))
    for material in MATERIALS:
        build_pdf(material, font_name)
        print(f"built {material['pdf']}")


if __name__ == "__main__":
    main()
